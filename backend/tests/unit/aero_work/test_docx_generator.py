"""DOCX 생성기 — Word 패키지와 문단 구조 단위 검증."""

from __future__ import annotations

import io

from docx import Document

from app.modules.aero_work.docx_generator import build_docx, build_docx_document


def _document(data: bytes) -> Document:
    return Document(io.BytesIO(data))


def test_build_docx_document_creates_readable_word_package() -> None:
    data = build_docx_document('문서 제목', ['제목', '본문 1', '본문 2'])

    assert data.startswith(b'PK\x03\x04')
    document = _document(data)
    assert [paragraph.text for paragraph in document.paragraphs] == ['제목', '본문 1', '본문 2']
    assert document.paragraphs[0].runs[0].bold is True
    assert document.core_properties.title == '문서 제목'


def test_build_docx_document_drops_blank_paragraphs_and_falls_back_to_title() -> None:
    document = _document(build_docx_document('  문서 제목  ', ['', '  ', ' 본문 ']))
    assert [paragraph.text for paragraph in document.paragraphs] == ['본문']

    empty_document = _document(build_docx_document('  문서 제목  ', ['', '  ']))
    assert [paragraph.text for paragraph in empty_document.paragraphs] == ['문서 제목']


def test_build_docx_splits_body_into_paragraphs() -> None:
    document = _document(build_docx('제목', '첫 문단\n\n 둘째 문단 '))
    assert [paragraph.text for paragraph in document.paragraphs] == ['제목', '첫 문단', '둘째 문단']
