"""본문 추출 — DOCX(python-docx)·HWPX(네이티브)·PDF(pypdf) 단위 검증."""

from __future__ import annotations

from pathlib import Path

from app.modules.aero_work.hwpx_generator import build_hwpx_document
from app.modules.aero_work.text_extract import extract_text, is_supported

_PDF_TEXT = 'Aero Work PDF'


def _make_pdf(text: str) -> bytes:
    stream = b'BT /F1 24 Tf 72 700 Td (' + text.encode('latin-1') + b') Tj ET'
    objects = [
        b'<</Type/Catalog/Pages 2 0 R>>',
        b'<</Type/Pages/Kids[3 0 R]/Count 1>>',
        b'<</Type/Page/Parent 2 0 R/MediaBox[0 0 612 792]'
        b'/Resources<</Font<</F1 5 0 R>>>>/Contents 4 0 R>>',
        b'<</Length ' + str(len(stream)).encode() + b'>>stream\n' + stream + b'\nendstream',
        b'<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>',
    ]
    pdf = b'%PDF-1.4\n'
    offsets: list[int] = []
    for index, obj in enumerate(objects, 1):
        offsets.append(len(pdf))
        pdf += str(index).encode() + b' 0 obj' + obj + b'endobj\n'
    xref_pos = len(pdf)
    pdf += b'xref\n0 ' + str(len(objects) + 1).encode() + b'\n0000000000 65535 f \n'
    for offset in offsets:
        pdf += ('%010d 00000 n \n' % offset).encode()
    pdf += b'trailer<</Size ' + str(len(objects) + 1).encode() + b'/Root 1 0 R>>\n'
    pdf += b'startxref\n' + str(xref_pos).encode() + b'\n%%EOF'
    return pdf


def test_supported_suffixes_include_office() -> None:
    assert is_supported(Path('a.pdf'))
    assert is_supported(Path('a.docx'))
    assert is_supported(Path('a.hwpx'))
    assert not is_supported(Path('a.hwp'))  # 구 바이너리 HWP 는 후속


def test_extract_docx(tmp_path: Path) -> None:
    import docx

    document = docx.Document()
    document.add_paragraph('출장 정산 규정 문단')
    document.add_paragraph('두 번째 문단입니다')
    path = tmp_path / 'sample.docx'
    document.save(str(path))

    text = extract_text(path)
    assert '출장 정산 규정 문단' in text
    assert '두 번째 문단입니다' in text


def test_extract_hwpx(tmp_path: Path) -> None:
    path = tmp_path / 'sample.hwpx'
    path.write_bytes(build_hwpx_document('보고서', ['보고서', '본문 첫 줄', '본문 둘째 줄']))

    text = extract_text(path)
    assert '보고서' in text
    assert '본문 첫 줄' in text
    assert '본문 둘째 줄' in text


def test_extract_pdf(tmp_path: Path) -> None:
    path = tmp_path / 'sample.pdf'
    path.write_bytes(_make_pdf(_PDF_TEXT))
    text = extract_text(path)
    assert _PDF_TEXT in text


def test_corrupt_binary_returns_empty(tmp_path: Path) -> None:
    path = tmp_path / 'broken.pdf'
    path.write_bytes(b'not a real pdf')
    assert extract_text(path) == ''  # best-effort: 손상 파일은 빈 문자열
