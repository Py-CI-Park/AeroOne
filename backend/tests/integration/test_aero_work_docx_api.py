"""Aero Work DOCX REST — 인증·CSRF·양식·승인본 다운로드 통합 검증."""

from __future__ import annotations

import io
import zipfile

from docx import Document


DOCX_MEDIA_TYPE = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'


def _document(data: bytes) -> Document:
    return Document(io.BytesIO(data))


def test_docx_document_anonymous_rejected(client) -> None:
    response = client.post('/api/v1/aero-work/document/docx', json={'title': 'x', 'body': 'y'})
    assert response.status_code == 401


def test_docx_document_requires_csrf(client) -> None:
    assert client.post('/api/v1/auth/login', json={'username': 'admin', 'password': 'password'}).status_code == 200
    response = client.post('/api/v1/aero-work/document/docx', json={'title': 'x', 'body': 'y'})
    assert response.status_code == 403


def test_docx_document_generates_word_package_and_records(csrf_client) -> None:
    response = csrf_client.post(
        '/api/v1/aero-work/document/docx',
        json={'title': '출장 결과 보고', 'body': '대전 방문함.\n특이사항 없음.'},
    )

    assert response.status_code == 200, response.text
    assert response.headers['content-type'] == DOCX_MEDIA_TYPE
    assert 'attachment' in response.headers['content-disposition']
    assert response.content.startswith(b'PK\x03\x04')
    paragraphs = [paragraph.text for paragraph in _document(response.content).paragraphs]
    assert paragraphs[0] == '출장 결과 보고'
    assert '□ 요약: 대전 방문함.' in paragraphs
    assert '◦ 특이사항 없음.' in paragraphs

    activities = csrf_client.get('/api/v1/aero-work/activity').json()['activities']
    assert activities[0]['kind'] == 'document.generate'
    assert activities[0]['summary'] == 'DOCX 1페이지 보고서 생성 "출장 결과 보고"'


def test_docx_official_format_applies_hierarchy(csrf_client) -> None:
    response = csrf_client.post(
        '/api/v1/aero-work/document/docx',
        json={'title': '협조 요청', 'body': '자료 제출 협조\n기한 엄수', 'format': 'official'},
    )

    assert response.status_code == 200, response.text
    paragraphs = [paragraph.text for paragraph in _document(response.content).paragraphs]
    assert any(paragraph.startswith('수신') for paragraph in paragraphs)
    assert any(paragraph.startswith('제목') for paragraph in paragraphs)
    assert '1. 자료 제출 협조' in paragraphs
    assert any(paragraph.endswith('끝.') for paragraph in paragraphs)


def test_approved_saved_document_downloads_docx_or_default_hwpx(csrf_client) -> None:
    created = csrf_client.post(
        '/api/v1/aero-work/document/save-request',
        json={'title': '협조 요청', 'body': '자료 제출 협조', 'format': 'official'},
    )
    assert created.status_code == 201, created.text
    document_id = created.json()['id']
    assert csrf_client.post(f'/api/v1/aero-work/document/saved/{document_id}/approve').status_code == 200

    docx_response = csrf_client.get(f'/api/v1/aero-work/document/saved/{document_id}/download?kind=docx')
    assert docx_response.status_code == 200, docx_response.text
    assert docx_response.headers['content-type'] == DOCX_MEDIA_TYPE
    assert 'document.docx' in docx_response.headers['content-disposition']
    assert docx_response.content.startswith(b'PK\x03\x04')
    assert '1. 자료 제출 협조' in [paragraph.text for paragraph in _document(docx_response.content).paragraphs]

    hwpx_response = csrf_client.get(f'/api/v1/aero-work/document/saved/{document_id}/download')
    assert hwpx_response.status_code == 200, hwpx_response.text
    assert hwpx_response.headers['content-type'] == 'application/hwp+zip'
    assert 'document.hwpx' in hwpx_response.headers['content-disposition']
    section = zipfile.ZipFile(io.BytesIO(hwpx_response.content)).read('Contents/section0.xml').decode('utf-8')
    assert '1. 자료 제출 협조' in section
