"""DOCX 문서 생성 — 제목과 구조화된 문단을 Word 문서 바이트로 조립한다."""

from __future__ import annotations

import io

from docx import Document
from docx.shared import Pt

from app.modules.aero_work.hwpx_generator import split_paragraphs


def build_docx_document(title: str, paragraphs: list[str]) -> bytes:
    """양식별 구조화 문단을 DOCX로 조립한다. 첫 문단은 제목으로 강조한다."""
    title = (title or '무제').strip() or '무제'
    cleaned = [paragraph.strip() for paragraph in paragraphs if paragraph and paragraph.strip()]
    if not cleaned:
        cleaned = [title]
    document = Document()
    document.core_properties.title = title

    heading = document.add_paragraph(cleaned[0])
    heading.runs[0].bold = True
    heading.runs[0].font.size = Pt(16)
    for paragraph in cleaned[1:]:
        document.add_paragraph(paragraph)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def build_docx(title: str, body: str) -> bytes:
    """제목과 본문 문자열을 DOCX 문서 바이트로 조립한다."""
    return build_docx_document(title, [title, *split_paragraphs(body)])
