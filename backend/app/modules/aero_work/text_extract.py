"""지식폴더 파일 → 평문 텍스트 추출.

F4 로 텍스트/마크다운/HTML 에 더해 **PDF(pypdf) · DOCX(python-docx) · HWPX(네이티브 zip+xml)**
본문 추출을 지원한다. HWPX 는 OWPML(ZIP+XML)이라 별도 wheel 없이 zipfile+정규식으로 hp:t 텍스트를
뽑는다. 각 추출기는 best-effort — 손상/암호 파일은 빈 문자열을 돌려 색인이 중단되지 않게 한다.

HWP(구 5.0 바이너리) 본문 추출은 무거운 라이브러리가 필요해 후속 범위다.
"""

from __future__ import annotations

import html
import re
import zipfile
from pathlib import Path

from bs4 import BeautifulSoup

SUPPORTED_SUFFIXES = frozenset(
    {
        '.txt', '.md', '.markdown', '.mdx', '.html', '.htm', '.csv', '.tsv', '.log', '.rst', '.json',
        '.pdf', '.docx', '.hwpx',
    }
)
_HTML_SUFFIXES = frozenset({'.html', '.htm'})
_HWPX_SECTION_RE = re.compile(r'Contents/section\d+\.xml$', re.IGNORECASE)
_HWPX_TEXT_RE = re.compile(r'<hp:t[^>]*>(.*?)</hp:t>', re.DOTALL)


def is_supported(path: Path) -> bool:
    return path.suffix.lower() in SUPPORTED_SUFFIXES


def _extract_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        return '\n'.join((page.extract_text() or '') for page in reader.pages)
    except Exception:  # noqa: BLE001 — 손상/암호 PDF 는 빈 문자열로 강등
        return ''


def _extract_docx(path: Path) -> str:
    try:
        import docx

        document = docx.Document(str(path))
        parts = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                parts.append('\t'.join(cell.text for cell in row.cells))
        return '\n'.join(parts)
    except Exception:  # noqa: BLE001
        return ''


def _extract_hwpx(path: Path) -> str:
    try:
        chunks: list[str] = []
        with zipfile.ZipFile(path) as archive:
            for name in sorted(archive.namelist()):
                if not _HWPX_SECTION_RE.search(name):
                    continue
                xml = archive.read(name).decode('utf-8', errors='replace')
                for match in _HWPX_TEXT_RE.finditer(xml):
                    chunks.append(html.unescape(match.group(1)))
        return '\n'.join(chunks)
    except Exception:  # noqa: BLE001
        return ''


def extract_text(path: Path) -> str:
    """지원 파일에서 평문을 추출한다. 바이너리 계열(PDF/DOCX/HWPX)은 전용 추출기를 쓴다."""

    suffix = path.suffix.lower()
    if suffix == '.pdf':
        return _extract_pdf(path)
    if suffix == '.docx':
        return _extract_docx(path)
    if suffix == '.hwpx':
        return _extract_hwpx(path)

    raw = path.read_text(encoding='utf-8', errors='replace')
    if suffix in _HTML_SUFFIXES:
        soup = BeautifulSoup(raw, 'html.parser')
        for tag in soup(['script', 'style']):
            tag.decompose()
        return soup.get_text(separator='\n')
    return raw
